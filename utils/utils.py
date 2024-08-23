import platform
import re
import os
from dotenv import load_dotenv, find_dotenv
from typing import Literal, List, Union, Dict
from pydantic import BaseModel, Field
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

load_dotenv(find_dotenv())


class ChatMessage(BaseModel):
    role: Literal["user", "assistant", "system"]
    content: str


def get_os_type() -> Literal['Windows', 'macOS']:
    """
    判断当前操作系统是macOS还是Windows
    """
    os_type = platform.system()
    if os_type == "Darwin":
        return "macOS"
    elif os_type == "Windows":
        return "Windows"
    else:
        raise EnvironmentError("仅支持macOS和Windows操作系统。")


def transform_messages_type(history_messages: List[ChatMessage], has_system: bool = False):
    """ChatMessage类型转换为langchain ChatPrompt所需要的"""

    if not has_system:
        first_message = history_messages[0]
        # 如果存在system
        if first_message.role == 'system':
            history_messages.pop(0)

    chat_history = []
    for history_message in history_messages:
        if history_message.role == 'system':
            chat_history.append(SystemMessage(content=history_message.content))
        elif history_message.role == 'user':
            chat_history.append(HumanMessage(content=history_message.content))
        elif history_message.role == 'assistant':
            chat_history.append(AIMessage(content=history_message.content))
        else:
            raise ValueError('Please enter a valid role type(user or assistant or system).')

    return chat_history


# 智普模型
ZHIPU_MODEL = ['glm-4', 'glm-4-0520', 'glm-4-air']
# qwen模型
QWEN_MODEL = ['qwen-max', 'qwen-long', 'qwen-plus', 'qwen-turbo', 'qwen2-72b-instruct']


def change_model(model_name: str) -> Union[Dict[str, str], None]:
    """根据模型的name 切换模型"""

    pattern_glm = r'^glm\w+$'
    pattern_qwen = r'^qwen\w+$'

    if re.match(pattern_glm, model_name):
        # 匹配到glm-llm
        return {
            'model': model_name,
            'api_key': os.environ['ZHIPUAI_API_KEY'],
            'base_url': 'https://open.bigmodel.cn/api/paas/v4/'
        }

    if re.match(pattern_qwen, model_name):
        # 匹配到qwen-llm
        return {
            'model': '',
            'api_key': '',
            'base_url': ''
        }

    return None


def write_docx(doc_path: str, doc_content: str):
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    # from docx.shared import Cm, RGBColor, Pt

    doc = Document()

    doc.encoding = 'utf-8'

    """添加标题"""
    # para_head = doc.add_heading("标题", level=1)  # level代表标题级别
    # para_head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中

    """添加正文段落"""
    doc.add_paragraph(doc_content)

    """保存文档"""
    doc.save(doc_path)
